"""
文档生成模块 - 在内存中生成 .docx 并上传到 WPS 网盘
使用结构化 JSON + 预定义模板，支持多种文档格式。
流程：AI 生成 JSON → Python 渲染 → 上传 WPS
"""
import io
import json
import logging
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


async def generate_and_upload_document(
    access_token: str,
    title: str,
    content,
    dbsheet_file_id: Optional[str] = None,
    doc_type: str = "report",
    metadata: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    在内存中生成 Word 文档并上传到 WPS 网盘。

    Args:
        access_token: WPS OAuth token
        dbsheet_file_id: 多维表格 ID
        title: 文档标题（用于文件命名）
        content: JSON 字符串，描述文档结构
        doc_type: 文档类型 - official/report/notice/minutes/other
        metadata: 元数据（author/date/doc_number 等）

    content JSON 格式示例：
    {
      "sections": [
        {"type": "org_header", "text": "中国民航机场建设集团公司"},
        {"type": "doc_number", "text": "建设集团〔2026〕15号"},
        {"type": "title", "text": "标题"},
        {"type": "heading1", "text": "一、第一部分"},
        {"type": "paragraph", "text": "正文内容..."},
        {"type": "heading2", "text": "（一）第二层次"},
        {"type": "signature", "text": "中国民航机场建设集团公司"},
        {"type": "date", "text": "2026年4月7日"}
      ]
    }

    Returns:
        成功: {"ok": True, "file_name": "...", "cloud_file_id": "..."}
        失败: {"ok": False, "error": "..."}
    """
    if not title or not content:
        return {"ok": False, "error": "标题和内容不能为空"}

    try:
        # 解析文档结构：支持 dict（AI 直接传对象）和 str（旧格式兼容）
        if isinstance(content, dict):
            structure = content
            logger.info(f"[DOC] 接收到 dict，sections数量={len(structure.get('sections', []))}")
        else:
            print(f"[DOC] content type={type(content)}, len={len(content)}, first100={content[:100]}")
            try:
                structure = json.loads(content)
                print(f"[DOC] JSON解析成功，sections数量={len(structure.get('sections', []))}")
            except json.JSONDecodeError as _je:
                print(f"[DOC] JSON解析失败: {_je}")
                logger.error(f"JSON解析失败: {_je}\n原始content前300字符: {content[:300]}")

                # 尝试多种修复策略
                _fixed = content
                _strategies = []

                # 策略1：替换中文引号为标准ASCII引号
                _chn_left_quote = chr(8220)
                _chn_right_quote = chr(8221)
                if _chn_left_quote in _fixed or _chn_right_quote in _fixed:
                    _fixed = _fixed.replace(_chn_left_quote, '"').replace(_chn_right_quote, '"')
                    _strategies.append("替换中文引号")

                # 策略2：处理未转义的真实换行符
                import re as _re
                _fixed = _re.sub(r'("text":\s*)"([^"]*)\n([^"]*)"', r'\1"\2\\n\3"', _fixed)
                if '\\n' in _fixed and '\n' in content:
                    _strategies.append("处理未转义换行符")

                # 策略3：移除可能导致解析失败的控制字符
                _fixed = ''.join(c for c in _fixed if ord(c) >= 32 or c in '\n\t\r')

                # 策略4：尝试修复缺失的逗号
                _fixed = _re.sub(r'}\s*{', '},{', _fixed)
                if '},{' in _fixed:
                    _strategies.append("修复缺失逗号")

                _tried_strategies = []
                structure = None
                for _strategy_name in _strategies:
                    try:
                        structure = json.loads(_fixed)
                        print(f"[DOC] JSON修复成功: {_strategy_name}")
                        _tried_strategies.append(_strategy_name)
                        break
                    except json.JSONDecodeError:
                        _tried_strategies.append(f"{_strategy_name}(失败)")
                        continue

                if structure is None or "sections" not in structure:
                    print(f"[DOC] JSON所有修复都失败，退化为行拆分模式")
                    structure = {"sections": [{"type": "paragraph", "text": line} for line in content.split("\n") if line.strip()]}
                    logger.warning(f"JSON解析全部失败，已转换为普通文本模式。尝试的修复策略: {_tried_strategies}")
                else:
                    logger.info(f"JSON修复成功，使用策略: {_tried_strategies}")

        # 生成文档
        doc = Document()
        _apply_document_template(doc, structure, doc_type, metadata or {})

        # 序列化为字节
        buf = io.BytesIO()
        doc.save(buf)
        file_data = buf.getvalue()

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"{title}_{timestamp}.docx"
        logger.info(f"文档生成成功: {file_name}，大小: {len(file_data)} bytes")

        # 上传到 WPS 云盘
        if not access_token:
            return {"ok": False, "error": f"文档《{title}》已生成，但未连接 WPS，无法上传。"}

        if not dbsheet_file_id:
            return {"ok": False, "error": f"文档《{title}》已生成，但未指定 WPS 文件 ID。"}

        from agent.wps_client import upload_to_drive
        upload_result = await upload_to_drive(access_token, dbsheet_file_id, file_name, file_data,
                                               content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        cloud_file_id = upload_result.get("uploadId") or upload_result.get("id")
        if not cloud_file_id:
            raise ValueError(f"上传返回数据异常: {upload_result}")

        return {
            "ok": True,
            "file_name": file_name,
            "cloud_file_id": cloud_file_id,
            "message": f"文档《{title}》已生成并上传到 WPS\n文件名：{file_name}\n可在 WPS 云文档「AI附件」文件夹中查看"
        }

    except Exception as e:
        logger.error(f"文档生成/上传失败: {str(e)}", exc_info=True)
        return {"ok": False, "error": f"文档生成失败: {str(e)}"}


def _apply_document_template(doc: Document, structure: dict, doc_type: str, metadata: dict) -> None:
    """根据文档类型应用模板并渲染内容"""

    # 设置页面（所有文档类型通用）
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21.0)

    if doc_type == "official" or doc_type == "official_red":
        # 公文格式（红头/普通）：页边距特殊
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.7)
    else:
        # 普通文档：标准页边距
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 渲染各个 section
    sections = structure.get("sections", [])

    # 清理和过滤 sections：
    # 1. 普通文档（report/notice/minutes）不应该有公文要素（org_header/doc_number/red_line）
    # 2. 连续的 space 只保留一个
    # 3. title/heading 后面的 space 直接删除
    cleaned_sections = []
    prev_was_space = False
    prev_was_title_or_heading = False

    for sec in sections:
        sec_type = sec.get("type")

        # 过滤：普通文档不要公文要素
        if doc_type not in ("official", "official_red"):
            if sec_type in ("org_header", "doc_number", "red_line"):
                continue

        if sec_type == "space":
            # 如果前一个是 title 或 heading，跳过这个 space
            if prev_was_title_or_heading:
                prev_was_title_or_heading = False
                continue
            # 如果前一个是 space，跳过这个 space
            if not prev_was_space:
                cleaned_sections.append(sec)
                prev_was_space = True
        else:
            cleaned_sections.append(sec)
            prev_was_space = False
            prev_was_title_or_heading = (sec_type in ("title", "heading1", "heading2", "heading3"))

    # 通知类文档：没有 recipient 行时，在第一个 paragraph 前自动插入称谓行
    if doc_type in ("notice", "official", "official_red"):
        has_recipient = any(s.get("type") == "recipient" for s in cleaned_sections)
        if not has_recipient:
            for i, s in enumerate(cleaned_sections):
                if s.get("type") == "paragraph":
                    cleaned_sections.insert(i, {"type": "recipient", "text": "各部门、各位同事："})
                    break

    for sec in cleaned_sections:
        sec_type = sec.get("type", "paragraph")
        text = sec.get("text", "")

        # 自动修正单位名称
        if sec_type in ("org_header", "signature"):
            text = text.replace("中国民航机场建设集团公司", "民航机场规划设计研究总院有限公司")

        # 自动修正发文字号格式（部门名 → 单位简称）
        if sec_type == "doc_number":
            text = text.replace("科技质量部〔", "规划总院〔")
            text = text.replace("科技管理部〔", "规划总院〔")

        if sec_type == "org_header":
            _add_org_header(doc, text)
        elif sec_type == "doc_number":
            # 只有公文才渲染发文字号
            if doc_type in ("official", "official_red"):
                _add_doc_number(doc, text)
        elif sec_type == "red_line":
            _add_red_line(doc)
        elif sec_type == "title":
            _add_title(doc, text, doc_type)
        elif sec_type == "heading1":
            _add_heading1(doc, text)
        elif sec_type == "heading2":
            _add_heading2(doc, text)
        elif sec_type == "heading3":
            _add_heading3(doc, text)
        elif sec_type == "paragraph":
            _add_paragraph(doc, text)
        elif sec_type == "recipient":
            _add_paragraph(doc, text, no_indent=True)
        elif sec_type == "ending":
            _add_paragraph(doc, text, no_indent=True, space_before=True)
        elif sec_type == "signature":
            _add_signature(doc, text)
        elif sec_type == "date":
            _add_date(doc, text)


def _add_org_header(doc: Document, text: str) -> None:
    """发文机关标志：红色初号方正小标宋简体"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "方正小标宋简体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "方正小标宋简体")
    run.font.size = Pt(42)  # 初号
    run.font.color.rgb = RGBColor(255, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_doc_number(doc: Document, text: str) -> None:
    """发文字号：3号仿宋体，居中"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    run.font.size = Pt(16)  # 3号
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_red_line(doc: Document) -> None:
    """红色分隔线：1.5磅实线"""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt = 12/8
    bottom.set(qn('w:color'), 'FF0000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_title(doc: Document, text: str, doc_type: str) -> None:
    """标题：2号方正小标宋简体，不加粗，居中"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "方正小标宋简体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "方正小标宋简体")
    run.font.bold = False
    run.font.size = Pt(22)  # 2号
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)  # 段后0磅

    # 标题与正文之间空一行
    blank = doc.add_paragraph()
    blank.paragraph_format.space_after = Pt(0)
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.line_spacing = Pt(28)


def _add_heading1(doc: Document, text: str) -> None:
    """第一层次标题：3号黑体，段首空二字"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    run.font.size = Pt(16)  # 3号
    para.paragraph_format.first_line_indent = Pt(32)  # 空二字
    para.paragraph_format.line_spacing = Pt(28)  # 固定行距28磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅
    para.paragraph_format.space_before = Pt(0)  # 段前0磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅


def _add_heading2(doc: Document, text: str) -> None:
    """第二层次标题：3号楷体，段首空二字"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "楷体_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体_GB2312")
    run.font.size = Pt(16)
    para.paragraph_format.first_line_indent = Pt(32)
    para.paragraph_format.line_spacing = Pt(28)  # 固定行距28磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅
    para.paragraph_format.space_before = Pt(0)  # 段前0磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅


def _add_heading3(doc: Document, text: str) -> None:
    """第三层次标题：3号仿宋体加粗，段首空二字"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    run.font.size = Pt(16)
    run.font.bold = True
    para.paragraph_format.first_line_indent = Pt(32)
    para.paragraph_format.line_spacing = Pt(28)  # 固定行距28磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅
    para.paragraph_format.space_before = Pt(0)  # 段前0磅
    para.paragraph_format.space_after = Pt(0)  # 段后0磅


def _add_paragraph(doc: Document, text: str, no_indent: bool = False, space_before: bool = False) -> None:
    """正文段落：3号仿宋体，段首空二字，固定行距28磅
    no_indent=True 时强制顶格（用于 recipient/ending 类型）
    space_before=True 时段前空一行（用于结尾语）
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    run.font.size = Pt(16)
    if no_indent:
        para.paragraph_format.first_line_indent = Pt(0)
    else:
        para.paragraph_format.first_line_indent = Pt(32)
    para.paragraph_format.line_spacing = Pt(28)
    para.paragraph_format.space_before = Pt(28) if space_before else Pt(0)
    para.paragraph_format.space_after = Pt(0)


def _add_signature(doc: Document, text: str) -> None:
    """发文机关署名：3号仿宋体，右对齐，段前空一行"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    run.font.size = Pt(16)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(28)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(28)


def _add_date(doc: Document, text: str) -> None:
    """成文日期：3号仿宋体，右对齐"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    run.font.size = Pt(16)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(28)

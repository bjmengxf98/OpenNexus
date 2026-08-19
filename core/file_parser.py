"""
文件解析模块 - 支持 Word/PDF/图片/Markdown/纯文本
"""
import base64
import os
import tempfile
from pathlib import Path


def parse_file(file_path: str, filename: str, api_key: str = None,
               base_url: str = None, model: str = None,
               max_output_tokens: int = None) -> str:
    """
    解析文件，返回文本内容。
    图片类型需要多模态LLM，需传入api_key。
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md", ".csv"):
        return _parse_text(file_path)
    elif suffix in (".docx",):
        return _parse_docx(file_path)
    elif suffix in (".xlsx", ".xls"):
        return _parse_excel(file_path)
    elif suffix == ".pdf":
        return _parse_pdf(file_path, api_key, base_url, model, max_output_tokens)
    elif suffix in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"):
        if not api_key:
            return "[图片文件需要配置大模型API Key才能识别]"
        return _parse_image(file_path, api_key, base_url, model, max_output_tokens)
    else:
        return f"[不支持的文件类型: {suffix}]"


def _parse_text(file_path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312"):
        try:
            return Path(file_path).read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return "[文本文件编码无法识别]"


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Word文件解析失败: {e}]"


def _parse_excel(file_path: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"【工作表：{sheet_name}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                row_text = " | ".join(cells).strip(" |")
                if row_text.replace("|", "").strip():
                    parts.append(row_text)
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[解析Excel需要安装 openpyxl：pip install openpyxl]"
    except Exception as e:
        return f"[Excel解析失败: {e}]"


def _parse_pdf(file_path: str, api_key: str = None,
               base_url: str = None, model: str = None,
               max_output_tokens: int = None) -> str:
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)
        parts = []
        for page in doc:
            text = page.get_text().strip()
            if text:
                parts.append(text)
        doc.close()
        if parts:
            full_text = "\n".join(parts)
            # 提取内容过少（页码/水印/元数据残留），仍视为扫描件走视觉模型
            if len(full_text.strip()) >= 100:
                return full_text
        # 文字提取为空或内容极少 → 扫描件，尝试用视觉模型识别
        if api_key:
            return _parse_pdf_as_images(file_path, api_key, base_url, model, max_output_tokens)
        return "[PDF内容为空：该文件可能是扫描件/图片型PDF。请在设置中配置图片识别模型，系统将自动用视觉AI识别扫描件内容。]"
    except Exception as e:
        return f"[PDF解析失败: {e}]"


def _parse_pdf_as_images(file_path: str, api_key: str,
                          base_url: str = None, model: str = None,
                          max_output_tokens: int = None) -> str:
    """将PDF每页渲染为图片，调用视觉模型识别文字"""
    try:
        import fitz
        doc = fitz.open(file_path)
        results = []
        for page_num, page in enumerate(doc):
            mat = fitz.Matrix(2, 2)  # 2倍缩放，提高识别精度
            pix = page.get_pixmap(matrix=mat)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            pix.save(tmp.name)
            tmp.close()
            try:
                text = _parse_image(tmp.name, api_key, base_url, model, max_output_tokens)
                if text and not text.startswith("["):
                    results.append(f"[第{page_num + 1}页]\n{text}")
            finally:
                try:
                    os.unlink(tmp.name)
                except Exception:
                    pass
        doc.close()
        if results:
            return "\n\n".join(results)
        return "[PDF扫描件识别失败：视觉模型未返回内容]"
    except Exception as e:
        return f"[PDF转图片识别失败: {e}]"


def _parse_image(file_path: str, api_key: str, base_url: str = None,
                 model: str = None, max_output_tokens: int = None) -> str:
    """用多模态LLM识别图片内容"""
    import httpx
    import json

    suffix = Path(file_path).suffix.lower().lstrip(".")
    mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                "png": "image/png", "webp": "image/webp",
                "gif": "image/gif"}
    mime = mime_map.get(suffix, "image/png")

    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()

    # 如果没有传 base_url，根据 model 自动推断
    if not base_url:
        if model and "qwen" in model.lower():
            base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
        elif model and "doubao" in model.lower():
            base_url = "https://ark.cn-beijing.volces.com/api/v3"
        elif model and "gpt" in model.lower():
            base_url = "https://api.openai.com/v1"
        elif model and "gemini" in model.lower():
            base_url = "https://generativelanguage.googleapis.com/v1beta/openai"
        elif model and "claude" in model.lower():
            base_url = "https://api.anthropic.com/v1"
        else:
            base_url = "https://api.deepseek.com/v1"

    base_url = base_url.rstrip("/")
    model = model or "qwen-vl-plus"  # 默认用千问视觉模型

    payload = {
        "model": model,
        "max_tokens": max(128, min(int(max_output_tokens or 2000), 262_144)),
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url",
                 "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text",
                 "text": "请详细描述这张图片的全部文字内容，保持原有格式，不要遗漏任何信息。"}
            ]
        }]
    }

    try:
        resp = httpx.post(
            f"{base_url}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}",
                     "Content-Type": "application/json"},
            json=payload,
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[图片识别失败: {e}]"
